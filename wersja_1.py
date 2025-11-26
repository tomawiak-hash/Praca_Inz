import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
import os
import re
import json
from docxtpl import DocxTemplate
from io import BytesIO
import datetime
from docx import Document

# ----- Konfiguracja Aplikacji
st.set_page_config(page_title="Inteligentny Generator Szkoleń BHP", page_icon="🎓", layout="wide")

# Konfiguracja API
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
MODEL_NAME = 'gemini-1.5-flash' # Używamy nowszego, szybszego modelu

# ----- Inicjalizacja "pamięci" aplikacji (session_state)
if 'etap' not in st.session_state:
    st.session_state.etap = 1
if 'finalna_tresc' not in st.session_state:
    st.session_state.finalna_tresc = ""
if 'zapisana_firma' not in st.session_state:
    st.session_state.zapisana_firma = ""
if 'wybrany_zawod' not in st.session_state:
    st.session_state.wybrany_zawod = ""
if 'opis_zawodu' not in st.session_state:
    st.session_state.opis_zawodu = ""
if 'spis_tresci_do_tematyki' not in st.session_state:
    st.session_state.spis_tresci_do_tematyki = []
if 'cel_szkolenia_text' not in st.session_state:
    st.session_state.cel_szkolenia_text = ""
if 'tematyka_z_godzinami' not in st.session_state:
    st.session_state.tematyka_z_godzinami = []
if 'cached_test_content' not in st.session_state:
    st.session_state.cached_test_content = None
if 'cached_key_content' not in st.session_state:
    st.session_state.cached_key_content = None

# ----- Funkcje Narzędziowe (Helpers)
def generuj_dokument_z_tabela(nazwa_szablonu, context, dane_tabeli=None, mapowanie_kolumn=None, index_tabeli=0):
    """
    Uniwersalna funkcja do generowania dokumentów Word z dynamiczną tabelą.
    
    Args:
        nazwa_szablonu (str): Ścieżka do pliku .docx.
        context (dict): Słownik zmiennych do podmienienia w tekście (Jinja2).
        dane_tabeli (list): Lista słowników z danymi do wierszy tabeli.
        mapowanie_kolumn (list): Lista kluczy słownika odpowiadająca kolumnom (pomijając Lp.).
        index_tabeli (int): Którą tabelę w dokumencie wypełnić (domyślnie pierwszą [0]).
    """
    try:
        # Krok 1: Renderowanie zmiennych prostych (nagłówki, daty itp.)
        doc_tpl = DocxTemplate(nazwa_szablonu)
        doc_tpl.render(context)
        
        temp_bio = BytesIO()
        doc_tpl.save(temp_bio)
        temp_bio.seek(0)

        # Krok 2: Wypełnianie tabeli (jeśli podano dane)
        if dane_tabeli and mapowanie_kolumn:
            doc = Document(temp_bio)
            if doc.tables and len(doc.tables) > index_tabeli:
                table = doc.tables[index_tabeli]
                
                for i, wiersz_dane in enumerate(dane_tabeli):
                    row_cells = table.add_row().cells
                    
                    # Kolumna 0 to zawsze Lp.
                    row_cells[0].text = str(i + 1)
                    
                    # Reszta kolumn wg mapowania
                    for col_idx, klucz in enumerate(mapowanie_kolumn):
                        target_idx = col_idx + 1
                        if target_idx < len(row_cells):
                            wartosc = str(wiersz_dane.get(klucz, ''))
                            row_cells[target_idx].text = wartosc
            else:
                return None, f"Brak tabeli o indeksie {index_tabeli} w szablonie."

            final_bio = BytesIO()
            doc.save(final_bio)
            final_bio.seek(0)
            return final_bio, None
        
        # Jeśli brak danych tabeli, zwracamy wyrenderowany szablon
        return temp_bio, None

    except Exception as e:
        return None, str(e)

def generuj_docx_prosty(nazwa_szablonu, kontekst, nazwa_pliku_wynikowego):
    """Wrapper dla prostych dokumentów bez dynamicznych tabel."""
    try:
        doc = DocxTemplate(nazwa_szablonu)
        doc.render(kontekst)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
    except Exception as e:
        st.error(f"Błąd generowania pliku '{nazwa_pliku_wynikowego}': {e}")
        return None

# ----- Funkcje Logiki Biznesowej
def wczytaj_liste_zawodow_lokalnie():
    lista_zawodow = {
        "Administrator baz danych (252101)": "252101",
        "Specjalista administracji publicznej (242217)": "242217",
        "Specjalista do spraw kadr (242307)": "242307",
        "Kierownik biura (334101)": "334101",
        "Asystent dyrektora (334302)": "334302"
    }
    return lista_zawodow

@st.cache_data
def pobierz_opis_zawodu_lokalnie(kod_zawodu):
    sciezka_pliku = os.path.join('baza_zawodow', f'{kod_zawodu}.pdf')
    try:
        pelny_tekst = ""
        with open(sciezka_pliku, "rb") as f:
            pdf_reader = PdfReader(f)
            for page in pdf_reader.pages:
                pelny_tekst += (page.extract_text() or "") + "\n"
        return pelny_tekst
    except FileNotFoundError:
        return f"Błąd: Brak pliku {kod_zawodu}.pdf w folderze 'baza_zawodow'."
    except Exception as e:
        return f"Błąd odczytu pliku PDF {kod_zawodu}.pdf: {e}"

@st.cache_data
def laduj_baze_wiedzy(folder_path='baza_wiedzy'):
    pelny_tekst = ""
    if not os.path.isdir(folder_path):
        return "" # Ciche pominięcie, jeśli folder nie istnieje
    for nazwa_pliku in os.listdir(folder_path):
        sciezka_pliku = os.path.join(folder_path, nazwa_pliku)
        try:
            if nazwa_pliku.lower().endswith('.pdf'):
                with open(sciezka_pliku, "rb") as f:
                    pdf_reader = PdfReader(f)
                    if pdf_reader.is_encrypted: continue
                    for page in pdf_reader.pages:
                         pelny_tekst += (page.extract_text() or "") + "\n\n"
            elif nazwa_pliku.lower().endswith('.txt'):
                with open(sciezka_pliku, "r", encoding="utf-8") as f:
                    pelny_tekst += f.read() + "\n\n"
        except Exception as e:
            print(f"Błąd pliku {nazwa_pliku}: {e}")
    return pelny_tekst

def generuj_kompletne_szkolenie(firma, nazwa_zawodu, opis_zawodu, dodatkowe_zagrozenia):
    model = genai.GenerativeModel(MODEL_NAME)
    
    prompt = f"""
    Jesteś metodykiem BHP. Stwórz materiał szkoleniowy dla stanowiska '{nazwa_zawodu}' w firmie '{firma}'.
    
    STRUKTURA (BEZWZGLĘDNA):
    CZĘŚĆ 1: INSTRUKTAŻ OGÓLNY (11 punktów zgodnych z rozp. MGiP)
    CZĘŚĆ 2: INSTRUKTAŻ STANOWISKOWY (5 punktów - instruktaż, próbne wykonanie, praca samodzielna)

    WYTYCZNE:
    - Personalizuj treść w oparciu o OPIS ZAWODU i DODATKOWE ZAGROŻENIA poniżej.
    - Używaj Markdown (#, ##, ###).
    
    --- OPIS ZAWODU ---
    {opis_zawodu}
    --- DODATKOWE ZAGROŻENIA ---
    {dodatkowe_zagrozenia}
    """
    
    try:
        response = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(temperature=0.4))
        # Wyciągamy spis treści (linie zaczynające się od cyfry i kropki)
        st.session_state.spis_tresci_do_tematyki = re.findall(r"^(?:\d+)\.\s.*", response.text, re.MULTILINE)
        return response.text
    except Exception as e:
        st.error(f"Błąd API: {e}")
        return "Błąd generowania treści."

@st.cache_data
def generuj_cel_szkolenia(nazwa_szkolenia):
    try:
        model = genai.GenerativeModel(MODEL_NAME)
        prompt = f"Napisz jednozdaniowy, formalny cel szkolenia wstępnego BHP dla: '{nazwa_szkolenia}'."
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception:
        return "Zapoznanie pracownika z zagrożeniami i zasadami bezpiecznej pracy."

@st.cache_data
def generuj_test_bhp(_finalna_tresc):
    model = genai.GenerativeModel(MODEL_NAME)
    prompt = f"""
    Stwórz test BHP (10 pytań A,B,C) oraz klucz odpowiedzi na podstawie poniższego tekstu.
    
    FORMAT:
    1. Pytanie...
       A. ...
       B. ...
       C. ...
    (po 10 pytaniach wstaw linię): ---KLUCZ---
    1. A
    2. B
    ...

    TEKST:
    {_finalna_tresc[:30000]} 
    """
    try:
        response = model.generate_content(prompt)
        if "---KLUCZ---" in response.text:
            tresc_testu, klucz_odpowiedzi = response.text.split("---KLUCZ---", 1)
            return tresc_testu.strip(), klucz_odpowiedzi.strip()
        else:
            return response.text, None
    except Exception as e:
        st.error(f"Błąd generowania testu: {e}")
        return "Błąd.", None

@st.cache_data
def przypisz_godziny_do_tematow(_spis_tresci_lista):
    """Używa trybu JSON dla większej niezawodności."""
    model = genai.GenerativeModel(MODEL_NAME)
    tekst_spisu = "\n".join(_spis_tresci_lista)
    
    prompt = f"""
    Oszacuj godziny lekcyjne (45min) dla tematów BHP.
    Zwróć wynik WYŁĄCZNIE jako listę JSON:
    [
        {{"nazwa": "Temat 1", "godziny": 2}},
        {{"nazwa": "Temat 2", "godziny": 1}}
    ]
    Spis:
    {tekst_spisu}
    """
    
    try:
        response = model.generate_content(prompt)
        text_response = response.text.strip()
        # Czyszczenie markdowna json jeśli się pojawi
        if text_response.startswith("```json"):
            text_response = text_response[7:-3]
        elif text_response.startswith("```"):
             text_response = text_response[3:-3]

        return json.loads(text_response)
    except Exception as e:
        st.warning(f"Problem z automatycznym przypisaniem godzin: {e}")
        return []

# ----- Główny interfejs aplikacji
st.title("🎓 Inteligentny Generator Szkoleń BHP")

# --- DIAGNOSTYKA TABEL (Wklej tymczasowo) ---
if st.button("🕵️ Pokaż struktury tabel w szablonie Protokołu"):
    try:
        doc = Document("protokol_egzaminu_szablon_uproszczony.docx") # Upewnij się, że nazwa pliku jest poprawna
        st.write(f"Znaleziono łącznie {len(doc.tables)} tabel.")
        
        for i, table in enumerate(doc.tables):
            # Pobieramy tekst z pierwszego wiersza każdej tabeli
            if len(table.rows) > 0:
                cells_text = [cell.text.strip() for cell in table.rows[0].cells]
                st.info(f"🟦 Tabela indeks {i} | Liczba kolumn: {len(table.columns)}")
                st.code(f"Zawartość nagłówka: {cells_text}")
            else:
                st.warning(f"Tabela indeks {i} jest pusta.")
    except Exception as e:
        st.error(f"Błąd odczytu pliku: {e}")
# --- KONIEC DIAGNOSTYKI ---

# --- Etap 1: Wybór zawodu i generowanie treści ---
if st.session_state.etap == 1:
    st.header("Krok 1: Wybierz zawód i wygeneruj kompletne szkolenie")
    
    lista_zawodow = wczytaj_liste_zawodow_lokalnie()
    
    wybrany_zawod_nazwa = st.selectbox("Wybierz zawód z listy:", options=list(lista_zawodow.keys()), index=None, placeholder="Wybierz zawód...")
    dodatkowe_zagrozenia = st.text_area("Dodatkowe zagrożenia (opcjonalnie):", key="extra_hazards")
    nazwa_firmy = st.text_input("Nazwa firmy:", key="firma_input", value="Przykładowa Firma S.A.")
    
    if st.button("🚀 Generuj kompletne szkolenie"):
        if not wybrany_zawod_nazwa:
            st.warning("Proszę wybrać zawód z listy.")
        else:
            with st.spinner(f"Tworzenie materiałów dla: {wybrany_zawod_nazwa}..."):
                kod_zawodu = lista_zawodow[wybrany_zawod_nazwa]
                opis_zawodu = pobierz_opis_zawodu_lokalnie(kod_zawodu)
                
                if "Błąd:" in opis_zawodu:
                    st.error(opis_zawodu)
                else:
                    # Generowanie treści głównej
                    finalna_tresc = generuj_kompletne_szkolenie(nazwa_firmy, wybrany_zawod_nazwa, opis_zawodu, dodatkowe_zagrozenia)
                    
                    if "Błąd" not in finalna_tresc:
                        st.session_state.finalna_tresc = finalna_tresc
                        st.session_state.zapisana_firma = nazwa_firmy or "Firma"
                        st.session_state.wybrany_zawod = wybrany_zawod_nazwa
                        
                        # Generowanie metadanych w tle
                        st.session_state.cel_szkolenia_text = generuj_cel_szkolenia(f"Szkolenie BHP: {wybrany_zawod_nazwa}")
                        
                        if st.session_state.spis_tresci_do_tematyki:
                            st.session_state.tematyka_z_godzinami = przypisz_godziny_do_tematow(st.session_state.spis_tresci_do_tematyki)
                        else:
                            st.session_state.tematyka_z_godzinami = []

                        st.session_state.etap = 2
                        st.rerun()

# --- Etap 2: Weryfikacja i przejście do dokumentacji ---
elif st.session_state.etap == 2:
    st.header("✅ Krok 2: Weryfikacja treści")
    st.success("Szkolenie wygenerowane pomyślnie!")

    with st.expander("Podgląd treści szkolenia"):
        st.markdown(st.session_state.finalna_tresc)

    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            label="📥 Pobierz treść (.txt)",
            data=st.session_state.finalna_tresc.encode('utf-8'),
            file_name=f"Szkolenie_{st.session_state.wybrany_zawod}.txt",
            mime="text/plain"
        )
    with col2:
        if st.button("📄 Przejdź do generatora dokumentów"):
            st.session_state.etap = 3
            st.rerun()
            
    if st.button("🔙 Zacznij od nowa"):
        st.session_state.etap = 1
        st.rerun()

# --- Etap 3: Generator Dokumentacji ---
elif st.session_state.etap == 3:
    st.header("✅ Krok 3: Generator Dokumentacji")
    
    st.subheader("Dane wspólne")
    st.caption("Lista uczestników: Imię Nazwisko, Miejsce Pracy, Funkcja, Data Urodzenia (DD.MM.RRRR)")
    
    uczestnicy_input = st.text_area("Uczestnicy (jeden na linię):", height=100, key="uczestnicy_lista_input")
    
    # Parsowanie uczestników
    uczestnicy_dane_lista = []
    if uczestnicy_input:
        for i, linia in enumerate(uczestnicy_input.strip().splitlines()):
            czesci = [c.strip() for c in linia.split(',')]
            if len(czesci) == 4:
                uczestnicy_dane_lista.append({
                    'index': i + 1, 'imie_nazwisko': czesci[0], 'miejsce_pracy': czesci[1],
                    'funkcja': czesci[2], 'data_urodzenia': czesci[3], 'ocena': '', 'uwagi': ''
                })

    col_d1, col_d2 = st.columns(2)
    with col_d1:
        data_start = st.date_input("Start:", value=datetime.date.today())
        nr_kursu = st.text_input("Nr kursu:", "01/BHP/2025")
        kierownik_kursu = st.text_input("Kierownik:", "Jan Kowalski")
    with col_d2:
        data_koniec = st.date_input("Koniec:", value=datetime.date.today())
        miejscowosc = st.text_input("Miejscowość:", "Warszawa")
        data_wystawienia = st.date_input("Data wystawienia:", value=datetime.date.today())

    st.markdown("---")

    # --- Generowanie Zaświadczenia ---
    with st.container(border=True):
        st.subheader("📄 Zaświadczenie")
        wybrany_uczestnik = st.selectbox("Dla kogo:", options=[u['imie_nazwisko'] for u in uczestnicy_dane_lista], index=None)
        
        if st.button("Generuj Zaświadczenie"):
            if wybrany_uczestnik:
                osoba = next((u for u in uczestnicy_dane_lista if u['imie_nazwisko'] == wybrany_uczestnik), None)
                context = {
                    'nazwa_organizatora_szkolenia': st.session_state.zapisana_firma,
                    'imie_nazwisko': osoba['imie_nazwisko'], 
                    'data_urodzenia': osoba['data_urodzenia'],
                    'nazwa_szkolenia': f"Szkolenie wstępne BHP: {st.session_state.wybrany_zawod}",
                    'forma_szkolenia': "kurs",
                    'nazwa_organizatora': st.session_state.zapisana_firma,
                    'dzien_rozpoczecia': data_start.strftime("%d.%m.%Y"), 
                    'dzien_zakonczenia': data_koniec.strftime("%d.%m.%Y"),
                    'cel_szkolenia': st.session_state.cel_szkolenia_text, 
                    'miejscowosc_szkolenia': miejscowosc,
                    'data_wystawienia_zaswiadczenia': data_wystawienia.strftime("%d.%m.%Y"),
                    'nr_zaswiadczenia_wg_rejestru': f"{nr_kursu}/{osoba['index']}"
                }
                plik = generuj_docx_prosty("certyfikat_szablon.docx", context, "Certyfikat.docx")
                if plik:
                    st.download_button("Pobierz Zaświadczenie", plik, f"Zaswiadczenie_{osoba['imie_nazwisko']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.warning("Wybierz uczestnika.")

# --- Generowanie Tematyki (POPRAWIONA WERSJA DLA 4 KOLUMN) ---
    with st.container(border=True):
        st.subheader("📋 Tematyka Szkolenia")
        if st.button("Generuj Tematykę"):
            tematyka = st.session_state.tematyka_z_godzinami
            if tematyka:
                # Obliczamy sumę godzin
                total_h = sum(t.get('godziny', 0) for t in tematyka if isinstance(t.get('godziny'), int))
                
                # Przygotowujemy dane do wyświetlenia
                tematyka_display = []
                for t in tematyka:
                    tematyka_display.append({
                        "nazwa": t.get('nazwa', ''),
                        "godziny": t.get('godziny', 0),
                        "praktyka": "0" # <--- DODANO: Wypełniamy kolumnę praktyki zerami
                    })

                # Dodajemy wiersz podsumowania
                tematyka_display.append({
                    "nazwa": "RAZEM:", 
                    "godziny": total_h,
                    "praktyka": "0"
                })

                # Generujemy dokument mapując 3 kolumny danych (plus Lp. które jest automatyczne)
                plik, blad = generuj_dokument_z_tabela(
                    "tematyka_szablon_uproszczony.docx", 
                    {}, 
                    tematyka_display, 
                    ['nazwa', 'godziny', 'praktyka'] # <--- ZMIANA: Mapujemy 3 kolumny danych
                )
                
                if plik:
                    st.download_button("Pobierz Tematykę", plik, "Tematyka.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.error(f"Błąd: {blad}")
            else:
                st.warning("Brak danych tematyki.")

    # --- Generowanie Wykazu Uczestników ---
    with st.container(border=True):
        st.subheader("👥 Wykaz Uczestników")
        if st.button("Generuj Wykaz"):
            if uczestnicy_dane_lista:
                plik, blad = generuj_dokument_z_tabela(
                    "wykaz_uczestnikow_szablon_uproszczony.docx",
                    {},
                    uczestnicy_dane_lista,
                    ['imie_nazwisko', 'miejsce_pracy', 'funkcja', 'data_urodzenia']
                )
                if plik:
                    st.download_button("Pobierz Wykaz", plik, "Wykaz_Uczestnikow.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.error(blad)
            else:
                st.warning("Brak uczestników.")

    # --- Generowanie Protokołu Egzaminu ---
    with st.container(border=True):
        st.subheader("📝 Protokół Egzaminu")
        komisja = st.text_area("Komisja (każda osoba w nowej linii):", "Przewodniczący\nCzłonek 1\nCzłonek 2")
        
        oceny_lista = []
        if uczestnicy_dane_lista:
            st.write("Oceny:")
            cols = st.columns(3)
            for i, u in enumerate(uczestnicy_dane_lista):
                with cols[i % 3]:
                    ocena = st.selectbox(f"{u['imie_nazwisko']}", ["5", "4", "3", "2"], key=f"ocena_{i}")
                    u['ocena'] = ocena # Aktualizujemy słownik uczestnika
        
        if st.button("Generuj Protokół"):
            komisja_arr = komisja.splitlines()
            context = {
                'rodzaj_szkolenia': f"Szkolenie BHP: {st.session_state.wybrany_zawod}",
                'data_egzaminu': data_koniec.strftime("%d.%m.%Y"),
                'nr_kursu': nr_kursu,
                'komisja_1_nazwisko': komisja_arr[0] if len(komisja_arr)>0 else "",
                'komisja_2_nazwisko': komisja_arr[1] if len(komisja_arr)>1 else "",
                'komisja_3_nazwisko': komisja_arr[2] if len(komisja_arr)>2 else "",
                'miejsce': miejscowosc,
                'nazwa_organizatora': st.session_state.zapisana_firma
            }
            
            # Uwaga: Protokół zazwyczaj ma tabelę jako drugą (index 1), bo pierwsza to np. nagłówek
            plik, blad = generuj_dokument_z_tabela(
                "protokol_egzaminu_szablon_uproszczony.docx",
                context,
                uczestnicy_dane_lista,
                ['imie_nazwisko', 'ocena', 'uwagi'], # Mapowanie kolumn
                index_tabeli=3 # Ważne!
            )
            if plik:
                 st.download_button("Pobierz Protokół", plik, "Protokol.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                 st.error(blad)

    # --- Generowanie Dziennika Zajęć ---
    with st.container(border=True):
        st.subheader("🗓️ Dziennik Zajęć")
        data_zajec = st.date_input("Data zajęć:", value=data_start)
        
        if st.button("Generuj Dziennik Zajęć"):
            tematyka = st.session_state.tematyka_z_godzinami
            if tematyka:
                # Przygotowanie danych: każdy temat to wiersz
                dane_do_dziennika = []
                for t in tematyka:
                    dane_do_dziennika.append({
                        'data': data_zajec.strftime("%d.%m.%Y"),
                        'godziny': t.get('godziny', 0),
                        'przedmiot': "Szkolenie BHP",
                        'temat': t.get('nazwa', '')
                    })
                
                context = {'nazwa_organizatora': st.session_state.zapisana_firma}
                plik, blad = generuj_dokument_z_tabela(
                    "dziennik_zajec_szablon_uproszczony.docx",
                    context,
                    dane_do_dziennika,
                    ['data', 'godziny', 'przedmiot', 'temat']
                )
                if plik:
                    st.download_button("Pobierz Dziennik Zajęć", plik, "Dziennik_Zajec.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.warning("Brak tematyki.")

    # --- Generowanie Dziennika Lekcyjnego ---
    with st.container(border=True):
        st.subheader("📓 Dziennik Lekcyjny")
        wykladowcy_raw = st.text_area("Wykładowcy (Imię Nazwisko, Miejsce, Funkcja):", "Jan Nowak, Firma X, Specjalista BHP")
        
        if st.button("Generuj Dziennik Lekcyjny"):
            # 1. Parsowanie wykładowców
            wykladowcy_lista = []
            for linia in wykladowcy_raw.strip().splitlines():
                parts = [p.strip() for p in linia.split(',', 2)]
                if len(parts) == 3:
                    wykladowcy_lista.append({
                        'imie_nazwisko': parts[0], 'miejsce_pracy': parts[1], 'funkcja': parts[2],
                        'przedmiot': '', 'godziny_plan': 0, 'godziny_wykonanie': 0
                    })
            
            if not wykladowcy_lista:
                st.error("Błąd formatu wykładowców.")
            else:
                # 2. Logika przydziału tematów
                tematyka = st.session_state.tematyka_z_godzinami
                total_plan = 0
                
                if tematyka:
                    for i, temat in enumerate(tematyka):
                        idx = i % len(wykladowcy_lista) # Rozdzielamy tematy cyklicznie
                        h = temat.get('godziny', 0)
                        if isinstance(h, int):
                            wykladowcy_lista[idx]['przedmiot'] += f"{temat.get('nazwa','')}\n"
                            wykladowcy_lista[idx]['godziny_plan'] += h
                            wykladowcy_lista[idx]['godziny_wykonanie'] += h
                            total_plan += h
                
                # Dodanie wiersza podsumowania
                wykladowcy_lista.append({
                    'imie_nazwisko': '', 'miejsce_pracy': '', 'funkcja': '', 
                    'przedmiot': 'RAZEM:', 'godziny_plan': total_plan, 'godziny_wykonanie': total_plan
                })

                context = {
                    'nazwa_organizatora': st.session_state.zapisana_firma,
                    'dla_kogo': f"Szkolenie dla: {st.session_state.wybrany_zawod}",
                    'data_od': data_start.strftime("%d.%m.%Y"), 'data_do': data_koniec.strftime("%d.%m.%Y"),
                    'miejsce': miejscowosc, 'kierownik_nazwisko': kierownik_kursu
                }

                plik, blad = generuj_dokument_z_tabela(
                    "dziennik_lekcyjny_szablon_uproszczony.docx",
                    context,
                    wykladowcy_lista,
                    ['imie_nazwisko', 'miejsce_pracy', 'funkcja', 'przedmiot', 'godziny_plan', 'godziny_wykonanie']
                )
                if plik:
                    st.download_button("Pobierz Dziennik Lekcyjny", plik, "Dziennik_Lekcyjny.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.error(blad)

    # --- Generowanie Rejestru Zaświadczeń ---
    with st.container(border=True):
        st.subheader("📚 Rejestr Zaświadczeń")
        if st.button("Generuj Rejestr"):
            rejestr_dane = []
            for i, u in enumerate(uczestnicy_dane_lista):
                rejestr_dane.append({
                    'numer': f"{nr_kursu}/{i+1}",
                    'imie_nazwisko': u['imie_nazwisko'],
                    'uwagi': ''
                })
            
            context = {
                'rodzaj_szkolenia': "wstępnego", 'nr_kursu': nr_kursu,
                'kierownik_nazwisko': kierownik_kursu,
                'data_wystawienia': data_wystawienia.strftime("%d.%m.%Y"),
                'nazwa_organizatora': st.session_state.zapisana_firma, 'miejsce': miejscowosc
            }
            
            # Rejestr: tabela z danymi to zazwyczaj tabela nr 2 (index 1)
            plik, blad = generuj_dokument_z_tabela(
                "rejestr_zaswiadczen_szablon_uproszczony.docx",
                context,
                rejestr_dane,
                ['numer', 'imie_nazwisko', 'podpis_dummy', 'uwagi'], # podpis_dummy to pusta kolumna
                index_tabeli=1
            )
            if plik:
                st.download_button("Pobierz Rejestr", plik, "Rejestr.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.error(blad)

    st.markdown("---")

    # --- Generowanie Testu ---
    with st.container(border=True):
        st.subheader("📝 Test Sprawdzający")
        
        if st.button("Generuj Test i Klucz"):
            with st.spinner("AI tworzy pytania..."):
                tresc, klucz = generuj_test_bhp(st.session_state.finalna_tresc)
                st.session_state.cached_test_content = tresc
                st.session_state.cached_key_content = klucz
        
        if st.session_state.cached_test_content:
            st.success("Test gotowy.")
            
            # Pobieranie Testu
            ctx_test = {
                'nazwa_szkolenia': f"Szkolenie: {st.session_state.wybrany_zawod}",
                'tresc_testu': st.session_state.cached_test_content
            }
            plik_test = generuj_docx_prosty("test_szablon.docx", ctx_test, "Test.docx")
            if plik_test:
                st.download_button("Pobierz Arkusz Testu", plik_test, "Test.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            # Pobieranie Klucza (jeśli istnieje)
            if st.session_state.cached_key_content:
                ctx_klucz = {'klucz_odpowiedzi': st.session_state.cached_key_content}
                plik_klucz = generuj_docx_prosty("klucz_odpowiedzi_szablon.docx", ctx_klucz, "Klucz.docx")
                if plik_klucz:
                    st.download_button("Pobierz Klucz Odpowiedzi", plik_klucz, "Klucz.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.warning("Klucz jest zintegrowany z treścią testu (AI nie rozdzieliło go poprawnie).")

    st.markdown("---")
    if st.button("🔄 Nowe Szkolenie"):
        st.session_state.etap = 1
        st.rerun()