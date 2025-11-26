import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
import os
import re
import json
from docxtpl import DocxTemplate
from io import BytesIO
import datetime
from docx import Document
import zipfile
import pandas as pd
import re
import time


# ----- Konfiguracja Aplikacji
st.set_page_config(page_title="Inteligentny Generator Szkoleń BHP", page_icon="🎓", layout="wide")


MODEL_NAME = 'gemini-3-pro-preview' # Używamy nowszego, szybszego modelu
try:
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
except Exception:
    pass

# ----- Inicjalizacja "pamięci" aplikacji (session_state)
if 'etap' not in st.session_state:
    st.session_state.etap = 1
if 'finalna_tresc' not in st.session_state:
    st.session_state.finalna_tresc = ""
if 'zapisana_firma' not in st.session_state:
    st.session_state.zapisana_firma = ""
if 'wybrany_zawod' not in st.session_state:
    st.session_state.wybrany_zawod = ""
if 'opis_zawodu' not in st.session_state:
    st.session_state.opis_zawodu = ""
if 'spis_tresci_do_tematyki' not in st.session_state:
    st.session_state.spis_tresci_do_tematyki = []
if 'cel_szkolenia_text' not in st.session_state:
    st.session_state.cel_szkolenia_text = ""
if 'tematyka_z_godzinami' not in st.session_state:
    st.session_state.tematyka_z_godzinami = []
if 'cached_test_content' not in st.session_state:
    st.session_state.cached_test_content = None
if 'cached_key_content' not in st.session_state:
    st.session_state.cached_key_content = None

# ----- Funkcje Narzędziowe (Helpers)
def generuj_dokument_z_tabela(nazwa_szablonu, context, dane_tabeli=None, mapowanie_kolumn=None, index_tabeli=0):
    """
    Uniwersalna funkcja do generowania dokumentów Word z dynamiczną tabelą.
    
    Args:
        nazwa_szablonu (str): Ścieżka do pliku .docx.
        context (dict): Słownik zmiennych do podmienienia w tekście (Jinja2).
        dane_tabeli (list): Lista słowników z danymi do wierszy tabeli.
        mapowanie_kolumn (list): Lista kluczy słownika odpowiadająca kolumnom (pomijając Lp.).
        index_tabeli (int): Którą tabelę w dokumencie wypełnić (domyślnie pierwszą [0]).
    """
    try:
        # Krok 1: Renderowanie zmiennych prostych (nagłówki, daty itp.)
        doc_tpl = DocxTemplate(nazwa_szablonu)
        doc_tpl.render(context)
        
        temp_bio = BytesIO()
        doc_tpl.save(temp_bio)
        temp_bio.seek(0)

        # Krok 2: Wypełnianie tabeli (jeśli podano dane)
        if dane_tabeli and mapowanie_kolumn:
            doc = Document(temp_bio)
            if doc.tables and len(doc.tables) > index_tabeli:
                table = doc.tables[index_tabeli]
                
                for i, wiersz_dane in enumerate(dane_tabeli):
                    row_cells = table.add_row().cells
                    
                    # Kolumna 0 to zawsze Lp.
                    row_cells[0].text = str(i + 1)
                    
                    # Reszta kolumn wg mapowania
                    for col_idx, klucz in enumerate(mapowanie_kolumn):
                        target_idx = col_idx + 1
                        if target_idx < len(row_cells):
                            wartosc = str(wiersz_dane.get(klucz, ''))
                            row_cells[target_idx].text = wartosc
            else:
                return None, f"Brak tabeli o indeksie {index_tabeli} w szablonie."

            final_bio = BytesIO()
            doc.save(final_bio)
            final_bio.seek(0)
            return final_bio, None
        
        # Jeśli brak danych tabeli, zwracamy wyrenderowany szablon
        return temp_bio, None

    except Exception as e:
        return None, str(e)

def generuj_docx_prosty(nazwa_szablonu, kontekst, nazwa_pliku_wynikowego):
    """Wrapper dla prostych dokumentów bez dynamicznych tabel."""
    try:
        doc = DocxTemplate(nazwa_szablonu)
        doc.render(kontekst)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
    except Exception as e:
        st.error(f"Błąd generowania pliku '{nazwa_pliku_wynikowego}': {e}")
        return None

# ----- Funkcje Logiki Biznesowej
def wczytaj_liste_zawodow_lokalnie():
    lista_zawodow = {
        "Administrator baz danych (252101)": "252101",
        "Specjalista administracji publicznej (242217)": "242217",
        "Specjalista do spraw kadr (242307)": "242307",
        "Kierownik biura (334101)": "334101",
        "Asystent dyrektora (334302)": "334302"
    }
    return lista_zawodow

@st.cache_data
def pobierz_opis_zawodu_lokalnie(kod_zawodu):
    sciezka_pliku = os.path.join('baza_zawodow', f'{kod_zawodu}.pdf')
    try:
        pelny_tekst = ""
        with open(sciezka_pliku, "rb") as f:
            pdf_reader = PdfReader(f)
            for page in pdf_reader.pages:
                pelny_tekst += (page.extract_text() or "") + "\n"
        return pelny_tekst
    except FileNotFoundError:
        return f"Błąd: Brak pliku {kod_zawodu}.pdf w folderze 'baza_zawodow'."
    except Exception as e:
        return f"Błąd odczytu pliku PDF {kod_zawodu}.pdf: {e}"

@st.cache_data
def laduj_baze_wiedzy(folder_path='baza_wiedzy'):
    pelny_tekst = ""
    if not os.path.isdir(folder_path):
        return "" # Ciche pominięcie, jeśli folder nie istnieje
    for nazwa_pliku in os.listdir(folder_path):
        sciezka_pliku = os.path.join(folder_path, nazwa_pliku)
        try:
            if nazwa_pliku.lower().endswith('.pdf'):
                with open(sciezka_pliku, "rb") as f:
                    pdf_reader = PdfReader(f)
                    if pdf_reader.is_encrypted: continue
                    for page in pdf_reader.pages:
                         pelny_tekst += (page.extract_text() or "") + "\n\n"
            elif nazwa_pliku.lower().endswith('.txt'):
                with open(sciezka_pliku, "r", encoding="utf-8") as f:
                    pelny_tekst += f.read() + "\n\n"
        except Exception as e:
            print(f"Błąd pliku {nazwa_pliku}: {e}")
    return pelny_tekst

def generuj_kompletne_szkolenie(firma, nazwa_zawodu, opis_zawodu, dodatkowe_zagrozenia):
    model = genai.GenerativeModel(MODEL_NAME)
    
    prompt = f"""
    Jesteś Starszym Inspektorem BHP. Opracuj Szczegółowy Program Szkolenia Wstępnego (Instruktaż Ogólny + Stanowiskowy) dla stanowiska: '{nazwa_zawodu}' w firmie '{firma}'.

    STYL I TON:
    - Język: Formalny, urzędowy, imperatywny (np. "Zabrania się...", "Pracownik ma obowiązek...").
    - Unikaj "lania wody" i ogólników o tym, że BHP jest ważne. Przejdź od razu do konkretów.
    - Skup się na specyfice zawodu: {nazwa_zawodu}.

    WYMAGANA STRUKTURA:
    
    # CZĘŚĆ 1: INSTRUKTAŻ OGÓLNY
    (Opracuj zagadnienia ogólne: prawo pracy, pierwsza pomoc, ppoż - krótko i węzłowato).

    # CZĘŚĆ 2: INSTRUKTAŻ STANOWISKOWY (To jest najważniejsza część!)
    (Tutaj musisz być bardzo szczegółowy. Wykorzystaj poniższy 'OPIS ZAWODU' i 'ZAGROŻENIA').
    Podziel tę część na podpunkty:
    A. Charakterystyka stanowiska i środowiska pracy.
    B. Omówienie zagrożeń (czynniki fizyczne, chemiczne, psychofizyczne).
    C. Dokładna instrukcja bezpiecznego wykonywania pracy (krok po kroku).
    D. Środki ochrony indywidualnej (co konkretnie pracownik musi ubrać).
    E. Postępowanie w sytuacjach awaryjnych specyficznych dla tego stanowiska.

    --- OPIS ZAWODU ---
    {opis_zawodu}
    
    --- DODATKOWE ZAGROŻENIA OD UŻYTKOWNIKA ---
    {dodatkowe_zagrozenia}
    
    Nie dodawaj żadnych wstępów typu "Oto plan". Zacznij od tytułu szkolenia.
    """
    
    try:
        response = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(temperature=0.3)) # Zmniejszamy temperaturę dla większej konkretności
        st.session_state.spis_tresci_do_tematyki = re.findall(r"^(?:\d+)\.\s.*", response.text, re.MULTILINE)
        return response.text
    except Exception as e:
        st.error(f"Błąd API: {e}")
        return "Błąd generowania treści."
    
@st.cache_data
def generuj_cel_szkolenia(nazwa_szkolenia):
    """
    Generuje krótki, czysty tekst celu szkolenia bez zbędnych dopisków i formatowania.
    """
    try:
        model = genai.GenerativeModel(MODEL_NAME)
        
        # Bardziej rygorystyczny prompt
        prompt = f"""
        Jesteś automatem bazodanowym. Twoim zadaniem jest wygenerowanie krótkiego wpisu do dokumentacji.
        
        Zadanie: Napisz cel szkolenia wstępnego BHP dla stanowiska: '{nazwa_szkolenia}'.
        
        RYGORYSTYCZNE ZASADY:
        1. Zwróć TYLKO jedno zdanie.
        2. NIE dodawaj żadnych wstępów typu "Oczywiście", "Oto propozycja", "W odpowiedzi na...".
        3. NIE używaj żadnego formatowania Markdown (absolutny zakaz używania gwiazdek **).
        4. Cel ma być krótki, zwięzły i oficjalny (max 15-20 słów).
        5. Zacznij bezpośrednio od słów: "Przygotowanie pracownika do..." lub "Zapoznanie pracownika z..."
        """
        
        response = model.generate_content(prompt)
        tekst = response.text
        
        # --- Czyszczenie (Safety Net) ---
        # Usuwamy gwiazdki (bold), płotki i inne znaki markdown
        tekst = tekst.replace('*', '').replace('#', '').replace('_', '')
        
        # Usuwamy ewentualne "gadatliwe" wstępy, jeśli AI mimo wszystko je doda
        zbedne_zwroty = ["Oczywiście", "oto propozycja", "Oto cel", ":", "\n"]
        for zwrot in zbedne_zwroty:
            tekst = tekst.replace(zwrot, ' ')
            
        # Usuwamy podwójne spacje, które mogły powstać przy czyszczeniu
        tekst = " ".join(tekst.split())
        
        return tekst.strip()
        
    except Exception as e:
        # Fallback w razie błędu
        return "Przygotowanie pracownika do bezpiecznego wykonywania pracy na powierzonym stanowisku."

def rozplanuj_zajecia(tematyka_lista, data_start):
    """
    Rozkłada tematy na kolejne dni robocze (pon-pt), przestrzegając limitu 8h/dzień.
    Zwraca listę tematów z przypisaną datą oraz faktyczną datę zakończenia.
    """
    
    harmonogram = []
    aktualna_data = data_start
    dzienne_godziny = 0
    MAX_H_DZIEN = 8 # Limit 8 godzin lekcyjnych na dzień

    for temat in tematyka_lista:
        godziny_tematu = temat.get('godziny', 0)
        
        if not isinstance(godziny_tematu, int) or godziny_tematu <= 0:
            continue # Pomijamy tematy bez godzin

        # 1. Sprawdzanie, czy aktualna_data jest weekendem
        while aktualna_data.weekday() >= 5: # 5 = Sobota, 6 = Niedziela
            aktualna_data += datetime.timedelta(days=1)
            
        # 2. Jeśli dodanie tematu przekroczy limit 8h
        if dzienne_godziny + godziny_tematu > MAX_H_DZIEN:
            # Przesuwamy cały temat na następny dzień roboczy
            aktualna_data += datetime.timedelta(days=1)
            
            # Ponownie sprawdzamy, czy następny dzień nie jest weekendem
            while aktualna_data.weekday() >= 5:
                aktualna_data += datetime.timedelta(days=1)
            
            # Resetujemy licznik godzin dla nowego dnia
            dzienne_godziny = 0
        
        # 3. Przypisanie tematu do bieżącej daty
        harmonogram.append({
            'data': aktualna_data.strftime("%d.%m.%Y"), 
            'godziny': godziny_tematu,
            'przedmiot': "Szkolenie BHP", 
            'temat': temat.get('nazwa', 'Brak tematu')
        })
        
        # 4. Aktualizacja godzin na dziś
        dzienne_godziny += godziny_tematu

    # Faktyczna data zakończenia to data ostatniego wpisu
    faktyczna_data_koniec = aktualna_data 
    
    # Jeśli ostatni dzień był zajęty, aktualna_data przeszła już na kolejny dzień.
    # Używamy daty z ostatniego wpisu w harmonogramie.
    if harmonogram:
         ostatni_wpis_data_str = harmonogram[-1]['data']
         faktyczna_data_koniec = datetime.datetime.strptime(ostatni_wpis_data_str, "%d.%m.%Y").date()

    return harmonogram, faktyczna_data_koniec

@st.cache_data
def generuj_test_bhp(_finalna_tresc):
    """
    Generuje ambitniejszy test BHP z wiarygodnymi dystraktorami.
    """
    model = genai.GenerativeModel(MODEL_NAME)
    prompt = f"""
    Jesteś egzaminatorem Państwowej Inspekcji Pracy. 
    Przygotuj test sprawdzający wiedzę (10 pytań zamkniętych A,B,C) na podstawie poniższego materiału.

    WYMAGANIA JAKOŚCIOWE:
    1. POZIOM TRUDNOŚCI: Średni/Wysoki. Unikaj pytań oczywistych.
    2. DYSTRAKTORY (Błędne odpowiedzi): Muszą brzmieć prawdopodobnie i logicznie (nie mogą być śmieszne ani absurdalne). Muszą wymagać wiedzy, by je odrzucić.
    3. KONTEKST: Pytania mają dotyczyć konkretnych procedur, liczb, zasad działania, a nie ogólników.
    4. FORMATOWANIE: 
       - Brak wstępów.
       - Od razu "1. Treść pytania...".
       - Po 10 pytaniach linia: ---KLUCZ---
       - Potem klucz: "1. A" itd.

    MATERIAŁ ŹRÓDŁOWY:
    {_finalna_tresc[:35000]} 
    """
    try:
        response = model.generate_content(prompt)
        tekst = response.text
        
        # Filtr czyszczący (ten sam co wcześniej, bo działa dobrze)
        smieci = ["Jasne,", "Oto test", "propozycja", "***", "---", "Witaj", "Oczywiście", "##"]
        if "1." in tekst:
            index_startu = tekst.find("1.")
            prefix = tekst[:index_startu]
            for smiec in smieci:
                if smiec in prefix:
                    tekst = tekst[index_startu:]
                    break
        tekst = tekst.replace("***", "").replace("##", "")

        if "---KLUCZ---" in tekst:
            tresc_testu, klucz_odpowiedzi = tekst.split("---KLUCZ---", 1)
            return tresc_testu.strip(), klucz_odpowiedzi.strip()
        else:
            return tekst.strip(), None
            
    except Exception as e:
        st.error(f"Błąd generowania testu: {e}")
        return "Nie udało się wygenerować testu.", None
    
@st.cache_data
def przypisz_godziny_do_tematow(_spis_tresci_lista):
    """
    Przypisuje godziny lekcyjne do KAŻDEGO DETALICZNEGO TEMATU z listy, 
    aby zachować dużą ilość pozycji w dokumentacji (jak w starej wersji).
    """
    model = genai.GenerativeModel(MODEL_NAME)
    tekst_spisu = "\n".join(_spis_tresci_lista)
    
    prompt = f"""
    Jesteś metodykiem BHP. Twoim zadaniem jest przypisanie godzin lekcyjnych (45 min) do KAŻDEGO z poniższych tematów.
    
    ZASADY:
    1. PRIORYTET: Nie grupuj tematów. Zostaw KAŻDY temat jako oddzielną pozycję w wyjściowej liście. Zachowaj maksymalną szczegółowość listy.
    2. Przypisz minimum 1 godzinę do każdego krótkiego tematu.
    3. Tematyka Instruktażu Ogólnego nie powinna przekraczać 4 godzin, ale zachowaj szczegółowość tematów cząstkowych.
    4. Tematyka Ratownicza (Pierwsza Pomoc, PPOŻ, Wypadki) powinna mieć najwięcej pozycji, z godzinami 1 lub 2 na każdą.

    Zwróć wynik WYŁĄCZNIE jako listę JSON, bez żadnego wstępu:
    [
        {{"nazwa": "Nazwa tematu", "godziny": 1}},
        {{"nazwa": "Kolejny temat", "godziny": 1}},
        {{...}}
    ]
    
    SZCZEGÓŁOWY SPIS TREŚCI DO ANALIZY:
    {tekst_spisu}
    """
    
    try:
        response = model.generate_content(prompt)
        text_response = response.text.strip()
        
        # Standardowe czyszczenie JSON
        if text_response.startswith("```json"):
            text_response = text_response[7:-3]
        elif text_response.startswith("```"):
             text_response = text_response[3:-3]

        return json.loads(text_response)
    except Exception as e:
        st.warning(f"Błąd parsowania JSON. Spróbuj wygenerować ponownie. Błąd: {e}")
        return []

# ----- Główny interfejs aplikacji
st.title("🎓 Inteligentny Generator Szkoleń BHP")

# --- Etap 1: Wybór zawodu i generowanie treści ---
if st.session_state.etap == 1:
    st.header("Krok 1: Wybierz zawód i wygeneruj kompletne szkolenie")
    
    lista_zawodow = wczytaj_liste_zawodow_lokalnie()
    
    wybrany_zawod_nazwa = st.selectbox("Wybierz zawód z listy:", options=list(lista_zawodow.keys()), index=None, placeholder="Wybierz zawód...")
    dodatkowe_zagrozenia = st.text_area("Dodatkowe zagrożenia (opcjonalnie):", key="extra_hazards")
    nazwa_firmy = st.text_input("Nazwa firmy:", key="firma_input", value="Przykładowa Firma S.A.")
    
    if st.button("🚀 Generuj kompletne szkolenie"):
        if not wybrany_zawod_nazwa:
            st.warning("Proszę wybrać zawód z listy.")
        else:
            with st.spinner(f"Tworzenie materiałów dla: {wybrany_zawod_nazwa}..."):
                kod_zawodu = lista_zawodow[wybrany_zawod_nazwa]
                opis_zawodu = pobierz_opis_zawodu_lokalnie(kod_zawodu)
                
                if "Błąd:" in opis_zawodu:
                    st.error(opis_zawodu)
                else:
                    # Generowanie treści głównej
                    finalna_tresc = generuj_kompletne_szkolenie(nazwa_firmy, wybrany_zawod_nazwa, opis_zawodu, dodatkowe_zagrozenia)
                    
                    if "Błąd" not in finalna_tresc:
                        st.session_state.finalna_tresc = finalna_tresc
                        st.session_state.zapisana_firma = nazwa_firmy or "Firma"
                        st.session_state.wybrany_zawod = wybrany_zawod_nazwa
                        
                        # Generowanie metadanych w tle
                        st.session_state.cel_szkolenia_text = generuj_cel_szkolenia(f"Szkolenie BHP: {wybrany_zawod_nazwa}")
                        
                        if st.session_state.spis_tresci_do_tematyki:
                            st.session_state.tematyka_z_godzinami = przypisz_godziny_do_tematow(st.session_state.spis_tresci_do_tematyki)
                        else:
                            st.session_state.tematyka_z_godzinami = []

                        st.session_state.etap = 2
                        st.rerun()

# --- Etap 2: Weryfikacja i Edycja Programu ---
elif st.session_state.etap == 2:
    st.header("✅ Krok 2: Weryfikacja i Edycja Treści")
    st.success("Szkolenie wygenerowane pomyślnie!")

    # 1. EDYTOR HARMONOGRAMU (To jest Twoje nowe ulepszenie 2a)
    st.subheader("🛠️ Edytor Programu Szkolenia")
    st.info("Poniżej znajduje się wygenerowany program. Możesz **edytować nazwy**, **zmieniać godziny**, a także **dodawać i usuwać wiersze** przed wygenerowaniem dokumentów.")

    if st.session_state.tematyka_z_godzinami:
        # Konwersja listy słowników na DataFrame (tabelę)
        df = pd.DataFrame(st.session_state.tematyka_z_godzinami)
        
        # Konfiguracja wyświetlania kolumn
        column_config = {
            "nazwa": st.column_config.TextColumn(
                "Temat Szkolenia", 
                width="large", 
                required=True,
                help="Kliknij, aby edytować nazwę tematu"
            ),
            "godziny": st.column_config.NumberColumn(
                "Godziny (45min)", 
                min_value=1, 
                max_value=10, 
                step=1, 
                format="%d h",
                help="Liczba godzin lekcyjnych"
            )
        }

        # Wyświetlenie edytora
        # num_rows="dynamic" pozwala użytkownikowi dodawać i usuwać wiersze!
        edited_df = st.data_editor(
            df, 
            column_config=column_config, 
            use_container_width=True,
            num_rows="dynamic", 
            key="editor_tematyki",
            hide_index=True
        )

        # --- ZAPISYWANIE ZMIAN NA ŻYWO ---
        # Nadpisujemy stan aplikacji tym, co użytkownik zmienił w tabeli
        st.session_state.tematyka_z_godzinami = edited_df.to_dict('records')

        # Podsumowanie godzin na żywo
        total_h = edited_df['godziny'].sum()
        st.caption(f"📊 Łączna liczba godzin szkolenia: **{total_h}**")

    else:
        st.warning("Brak danych o tematyce. Spróbuj wygenerować szkolenie ponownie w Kroku 1.")

    st.markdown("---") 

    # 2. PODGLĄD TREŚCI TEKSTOWEJ (Ukryty w rozwijanym pasku, żeby nie zajmował miejsca)
    with st.expander("📖 Pokaż pełną treść merytoryczną szkolenia (Tekst)"):
        st.text_area("Edycja treści szkolenia (opcjonalnie):", value=st.session_state.finalna_tresc, height=300, key="edycja_tekstu_area")
        # Jeśli użytkownik zmieni tekst w tym polu, aktualizujemy go:
        st.session_state.finalna_tresc = st.session_state.edycja_tekstu_area

    st.markdown("---")
    
    # 3. PRZYCISKI NAWIGACJI
    col_btn1, col_btn2 = st.columns([1, 1])
    
    with col_btn1:
        st.download_button(
            label="📥 Pobierz samą treść (.txt)",
            data=st.session_state.finalna_tresc.encode('utf-8'),
            file_name=f"Szkolenie_{st.session_state.wybrany_zawod}.txt",
            mime="text/plain",
            use_container_width=True
        )

    with col_btn2:
        if st.button("📄 Zatwierdź i przejdź do dokumentów", type="primary", use_container_width=True):
            st.session_state.etap = 3
            st.rerun()
            
    if st.button("🔙 Wróć do wyboru zawodu", type="secondary"):
        st.session_state.etap = 1
        st.rerun()

# --- Etap 3: Generator Dokumentacji ---
elif st.session_state.etap == 3:
    st.header("✅ Krok 3: Generator Dokumentacji")
    
# --- SEKCJA DANYCH WSPÓLNYCH ---
    with st.container(border=True):
        st.subheader("🛠️ Konfiguracja danych")
        
        # 1. UCZESTNICY (Custom Label)
        st.markdown("**Lista uczestników** \n*Format: Imię Nazwisko, Miejsce Pracy, Funkcja, Data Urodzenia*", unsafe_allow_html=True)
        uczestnicy_input = st.text_area(
            label="Lista uczestników", # Etykieta dla systemów czytających (niewidoczna wizualnie)
            label_visibility="collapsed", # <--- UKRYWAMY STANDARDOWĄ ETYKIETĘ
            height=100, 
            key="uczestnicy_lista_input", 
            placeholder="Jan Kowalski, Biuro X, Księgowy, 12.05.1985\nAnna Nowak, Dział HR, Specjalista, 20.01.1990"
        )
        
        # Parsowanie i Walidacja Uczestników
        uczestnicy_dane_lista = []
        bledne_linie = []
        if uczestnicy_input:
            for i, linia in enumerate(uczestnicy_input.strip().splitlines()):
                linia = linia.strip()
                if not linia: continue
                czesci = [c.strip() for c in linia.split(',')]
                if len(czesci) == 4 and re.match(r"^\d{2}\.\d{2}\.\d{4}$", czesci[3]):
                    uczestnicy_dane_lista.append({'index': i+1, 'imie_nazwisko': czesci[0], 'miejsce_pracy': czesci[1], 'funkcja': czesci[2], 'data_urodzenia': czesci[3], 'ocena': '', 'uwagi': ''})
                else: bledne_linie.append(f"Błąd w linii {i+1}")

        if bledne_linie: st.error(f"Znaleziono błędy w {len(bledne_linie)} liniach.")
        if uczestnicy_dane_lista:
            with st.expander("🔍 Podgląd uczestników", expanded=False):
                st.dataframe(pd.DataFrame(uczestnicy_dane_lista)[['imie_nazwisko', 'miejsce_pracy', 'funkcja', 'data_urodzenia']], use_container_width=True, hide_index=True)

        st.markdown("---")

        # 2. BAZA KADRY (Custom Labels)
        st.markdown("### ⚙️ Baza Kadry")
        col_kadra1, col_kadra2 = st.columns(2)

        with col_kadra1:
            if 'baza_wykladowcow_text' not in st.session_state:
                st.session_state.baza_wykladowcow_text = "Jan Nowak, Firma BHP, Specjalista BHP\nAnna Kowalska, Firma Med, Ratownik"
            
            st.markdown("**1. Baza Wykładowców** \n*Format: Imię Nazwisko, Firma, Funkcja (do Dziennika Lekcyjnego)*", unsafe_allow_html=True)
            baza_wykladowcow = st.text_area(
                label="Baza Wykładowców",
                label_visibility="collapsed", # <--- UKRYWAMY STANDARDOWĄ ETYKIETĘ
                value=st.session_state.baza_wykladowcow_text,
                height=150,
                key="baza_wykladowcow_key"
            )
            opcje_wykladowcow = [x.strip() for x in baza_wykladowcow.splitlines() if x.strip()]

        with col_kadra2:
            if 'baza_komisji_text' not in st.session_state:
                st.session_state.baza_komisji_text = "Jan Nowak, Firma BHP, Przewodniczący\nAnna Kowalska, Firma BHP, Członek Komisji"
            
            st.markdown("**2. Baza Komisji Egzaminacyjnej** \n*Format: Imię Nazwisko, Firma, Funkcja (do Protokołu)*", unsafe_allow_html=True)
            baza_komisji = st.text_area(
                label="Baza Komisji",
                label_visibility="collapsed", # <--- UKRYWAMY STANDARDOWĄ ETYKIETĘ
                value=st.session_state.baza_komisji_text,
                height=150,
                key="baza_komisji_key"
            )
            opcje_komisji = [x.strip() for x in baza_komisji.splitlines() if x.strip()]

        st.markdown("---")

        # 3. HARMONOGRAM (Bez zmian)
        st.markdown("### 🗓️ Harmonogram Szkolenia")
        col_d1, col_d2 = st.columns(2)
        dzisiaj = datetime.date.today()
        
        with col_d1:
            data_start = st.date_input("Data rozpoczęcia:", key="doc_data_start", value=dzisiaj)
            nr_kursu = st.text_input("Numer kursu:", "01/BHP/2025", key="doc_nr_kursu")
            kierownik_kursu = st.text_input("Kierownik kursu:", "Anna Kowalska", key="doc_kierownik")
        
        sugerowana_data_koniec = data_start
        if st.session_state.tematyka_z_godzinami:
            _, obliczona_data = rozplanuj_zajecia(st.session_state.tematyka_z_godzinami, data_start)
            if obliczona_data > data_start: sugerowana_data_koniec = obliczona_data

        with col_d2:
            wartosc_domyslna = sugerowana_data_koniec
            if 'doc_data_koniec' in st.session_state:
                if st.session_state.doc_data_koniec >= data_start: wartosc_domyslna = st.session_state.doc_data_koniec
                else: wartosc_domyslna = data_start

            data_koniec = st.date_input("Data zakończenia:", key="doc_data_koniec", value=wartosc_domyslna, min_value=data_start)
            miejscowosc = st.text_input("Miejscowość:", "Łódź", key="doc_miejscowosc")
            
            wartosc_domyslna_wyst = data_koniec
            if 'doc_data_wyst' in st.session_state and st.session_state.doc_data_wyst >= data_koniec:
                 wartosc_domyslna_wyst = st.session_state.doc_data_wyst

            data_wystawienia = st.date_input("Data wystawienia dokumentów:", key="doc_data_wyst", value=wartosc_domyslna_wyst, min_value=data_koniec)

        st.session_state.faktyczna_data_koniec = data_koniec

    st.write("") # Odstęp

    # =========================================================
    # TU ZACZYNAJĄ SIĘ ZAKŁADKI (TABS) - GŁÓWNA ZMIANA
    # =========================================================
    
    tab1, tab2, tab3 = st.tabs(["📜 Zaświadczenia i Rejestr", "📅 Przebieg Szkolenia", "📝 Egzamin i Wyniki"])

    # --- ZAKŁADKA 1: ZAŚWIADCZENIA ---
    with tab1:
        st.info("Tutaj wygenerujesz dokumenty potwierdzające ukończenie szkolenia.")
        
        col_z1, col_z2 = st.columns(2)
        
        with col_z1:
            st.subheader("📄 Zaświadczenie Indywidualne")
            with st.container(border=True):
                wybrany_uczestnik = st.selectbox("Wybierz uczestnika:", options=[u['imie_nazwisko'] for u in uczestnicy_dane_lista], index=None)
                
                if st.button("Generuj Zaświadczenie", use_container_width=True):
                    if wybrany_uczestnik:
                        osoba = next((u for u in uczestnicy_dane_lista if u['imie_nazwisko'] == wybrany_uczestnik), None)
                        context = {
                            'nazwa_organizatora_szkolenia': st.session_state.zapisana_firma,
                            'imie_nazwisko': osoba['imie_nazwisko'], 
                            'data_urodzenia': osoba['data_urodzenia'],
                            'nazwa_szkolenia': f"Szkolenie wstępne BHP: {st.session_state.wybrany_zawod}",
                            'forma_szkolenia': "kurs",
                            'nazwa_organizatora': st.session_state.zapisana_firma,
                            'dzien_rozpoczecia': data_start.strftime("%d.%m.%Y"), 
                            'dzien_zakonczenia': data_koniec.strftime("%d.%m.%Y"),
                            'cel_szkolenia': st.session_state.cel_szkolenia_text, 
                            'miejscowosc_szkolenia': miejscowosc,
                            'data_wystawienia_zaswiadczenia': data_wystawienia.strftime("%d.%m.%Y"),
                            'nr_zaswiadczenia_wg_rejestru': f"{nr_kursu}/{osoba['index']}"
                        }
                        plik = generuj_docx_prosty("certyfikat_szablon.docx", context, "Certyfikat.docx")
                        if plik:
                            st.download_button("📥 Pobierz", plik, f"Zaswiadczenie_{osoba['imie_nazwisko']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                    else:
                        st.warning("Wybierz uczestnika z listy.")

        with col_z2:
            st.subheader("📚 Rejestr Wydanych Zaświadczeń")
            with st.container(border=True):
                st.write("Zbiorcza lista wszystkich wydanych zaświadczeń.")
                if st.button("Generuj Rejestr", use_container_width=True):
                    rejestr_dane = []
                    for i, u in enumerate(uczestnicy_dane_lista):
                        rejestr_dane.append({
                            'numer': f"{nr_kursu}/{i+1}",
                            'imie_nazwisko': u['imie_nazwisko'],
                            'uwagi': ''
                        })
                    
                    context = {
                        'rodzaj_szkolenia': "wstępnego", 'nr_kursu': nr_kursu,
                        'kierownik_nazwisko': kierownik_kursu,
                        'data_wystawienia': data_wystawienia.strftime("%d.%m.%Y"),
                        'nazwa_organizatora': st.session_state.zapisana_firma, 'miejsce': miejscowosc
                    }
                    
                    plik, blad = generuj_dokument_z_tabela(
                        "rejestr_zaswiadczen_szablon_uproszczony.docx",
                        context,
                        rejestr_dane,
                        ['numer', 'imie_nazwisko', 'podpis_dummy', 'uwagi'],
                        index_tabeli=2
                    )
                    if plik:
                        st.download_button("📥 Pobierz Rejestr", plik, "Rejestr.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                    else:
                        st.error(blad)

    # --- ZAKŁADKA 2: PRZEBIEG SZKOLENIA ---
    with tab2:
        st.info("Dokumentacja dotycząca programu, harmonogramu i realizacji zajęć.")
        
        col_p1, col_p2 = st.columns(2)
        
        with col_p1:
            st.subheader("📋 Tematyka Szkolenia")
            with st.container(border=True):
                if st.button("Generuj Tematykę", use_container_width=True):
                    # Mechanizm naprawczy (Plan B)
                    if not st.session_state.tematyka_z_godzinami and st.session_state.finalna_tresc:
                        with st.spinner("Odzyskiwanie tematów..."):
                            try:
                                model_fix = genai.GenerativeModel(MODEL_NAME)
                                prompt_fix = f"""
                                Przeanalizuj tekst i wypisz tematy z godzinami w JSON:
                                [ {{"nazwa": "Tytuł", "godziny": 1}} ]
                                TEKST: {st.session_state.finalna_tresc[:30000]}
                                """
                                response = model_fix.generate_content(prompt_fix)
                                text_resp = response.text.strip()
                                if text_resp.startswith("```json"): text_resp = text_resp[7:-3]
                                elif text_resp.startswith("```"): text_resp = text_resp[3:-3]
                                st.session_state.tematyka_z_godzinami = json.loads(text_resp)
                            except: pass
                    
                    # Generowanie
                    tematyka = st.session_state.tematyka_z_godzinami
                    if tematyka:
                        total_h = sum(t.get('godziny', 0) for t in tematyka if isinstance(t.get('godziny'), int))
                        tematyka_display = [{"nazwa": t.get('nazwa',''), "godziny": t.get('godziny',0), "praktyka": "0"} for t in tematyka]
                        tematyka_display.append({"nazwa": "RAZEM:", "godziny": total_h, "praktyka": "0"})

                        plik, blad = generuj_dokument_z_tabela("tematyka_szablon_uproszczony.docx", {}, tematyka_display, ['nazwa', 'godziny', 'praktyka'])
                        if plik:
                            st.download_button("📥 Pobierz Tematykę", plik, "Tematyka.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                        else:
                            st.error(blad)
                    else:
                        st.error("Brak danych tematyki.")

        with col_p2:
            st.subheader("🗓️ Dziennik Zajęć")
            with st.container(border=True):
                if st.button("Generuj Dziennik Zajęć", use_container_width=True):
                    tematyka = st.session_state.tematyka_z_godzinami
                    if tematyka:
                        zajecia, faktyczna_data = rozplanuj_zajecia(tematyka, data_start)
                        st.session_state.faktyczna_data_koniec = faktyczna_data
                        
                        doc_tpl = DocxTemplate("dziennik_zajec_szablon_uproszczony.docx")
                        doc_tpl.render({'nazwa_organizatora': st.session_state.zapisana_firma})
                        bio = BytesIO()
                        doc_tpl.save(bio)
                        bio.seek(0)
                        
                        doc = Document(bio)
                        if doc.tables:
                            table = doc.tables[0]
                            for i, z in enumerate(zajecia):
                                row = table.add_row().cells
                                if len(row) >= 6:
                                    row[0].text = str(i+1)
                                    row[1].text = z['data']
                                    row[2].text = str(z['godziny'])
                                    row[3].text = z['przedmiot']
                                    row[4].text = z['temat']
                        
                        final_bio = BytesIO()
                        doc.save(final_bio)
                        final_bio.seek(0)
                        st.download_button("📥 Pobierz Dziennik Zajęć", final_bio, "Dziennik_Zajec.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                    else:
                        st.error("Brak tematyki.")

        st.subheader("📓 Dziennik Lekcyjny")
        with st.container(border=True):
            st.write("Wybierz wykładowców, którzy prowadzili zajęcia (z listy w pasku bocznym).")
            
            # --- ZMIANA: MULTISELECT ZAMIAST TEXT_AREA ---
            wybrani_wykladowcy = st.multiselect(
                "Wybierz wykładowców do tego szkolenia:",
                options=opcje_wykladowcow,
                placeholder="Kliknij, aby wybrać osoby...",
                key="wykladowcy_multiselect"
            )
            
            if st.button("Generuj Dziennik Lekcyjny", use_container_width=True):
                # LOGIKA AUTOMATYCZNEGO WYBORU
                # Jeśli użytkownik nic nie wybrał, bierzemy wszystkich z bazy
                lista_do_przetworzenia = wybrani_wykladowcy if wybrani_wykladowcy else opcje_wykladowcow
                
                wykladowcy_lista = [] 
                
                if not lista_do_przetworzenia:
                    st.error("Brak wykładowców! Wpisz ich w Bazie Kadry (na górze strony).")
                else:
                    if not wybrani_wykladowcy:
                        st.info("ℹ️ Nie wybrano wykładowców ręcznie – użyto całej listy z bazy.")

                    bledy_formatu = False
                    for linia in lista_do_przetworzenia:
                        parts = [p.strip() for p in linia.split(',', 2)]
                        if len(parts) == 3:
                            wykladowcy_lista.append({
                                'imie_nazwisko': parts[0], 'miejsce_pracy': parts[1], 'funkcja': parts[2],
                                'przedmiot': 'Szkolenie wstępne BHP',
                                'godziny_plan': 0, 'godziny_wykonanie': 0
                            })
                        else:
                            st.error(f"Błąd formatu dla: '{linia}'. Wymagane 3 elementy oddzielone przecinkami.")
                            bledy_formatu = True
                    
                    if not bledy_formatu:
                        # ... Reszta kodu bez zmian ...
                        tematyka = st.session_state.tematyka_z_godzinami
                        total_plan = 0
                        
                        if tematyka:
                            for i, temat in enumerate(tematyka):
                                idx = i % len(wykladowcy_lista)
                                h = temat.get('godziny', 0)
                                if isinstance(h, int):
                                    wykladowcy_lista[idx]['godziny_plan'] += h
                                    wykladowcy_lista[idx]['godziny_wykonanie'] += h
                                    total_plan += h
                        
                        wykladowcy_lista.append({
                            'imie_nazwisko': '', 'miejsce_pracy': '', 'funkcja': '', 
                            'przedmiot': 'RAZEM:', 'godziny_plan': total_plan, 'godziny_wykonanie': total_plan
                        })

                        context = {
                            'nazwa_organizatora': st.session_state.zapisana_firma,
                            'dla_kogo': f"Szkolenie dla: {st.session_state.wybrany_zawod}",
                            'data_od': data_start.strftime("%d.%m.%Y"), 
                            'data_do': data_koniec.strftime("%d.%m.%Y"),
                            'miejsce': miejscowosc, 
                            'kierownik_nazwisko': kierownik_kursu,
                            'kierownik_miejsce_pracy_funkcja': "Kierownik Szkolenia" 
                        }

                        plik, blad = generuj_dokument_z_tabela(
                            "dziennik_lekcyjny_szablon_uproszczony.docx",
                            context,
                            wykladowcy_lista,
                            ['imie_nazwisko', 'miejsce_pracy', 'funkcja', 'przedmiot', 'godziny_plan', 'godziny_wykonanie'],
                            index_tabeli=4
                        )
                        
                        if plik:
                            st.download_button("📥 Pobierz Dziennik Lekcyjny", plik, "Dziennik_Lekcyjny.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                        else:
                            st.error(blad)

    # --- ZAKŁADKA 3: EGZAMIN ---
    with tab3:
        st.info("Dokumentacja związana z weryfikacją wiedzy uczestników.")
        
        # Dwie kolumny główne
        col_left, col_right = st.columns(2)
        
        # --- LEWA KOLUMNA: Wykaz i Test ---
        with col_left:
            st.subheader("👥 Wykaz Uczestników")
            with st.container(border=True):
                if st.button("Generuj Wykaz", use_container_width=True):
                    if uczestnicy_dane_lista:
                        plik, blad = generuj_dokument_z_tabela("wykaz_uczestnikow_szablon_uproszczony.docx", {}, uczestnicy_dane_lista, ['imie_nazwisko', 'miejsce_pracy', 'funkcja', 'data_urodzenia'])
                        if plik:
                            st.download_button("📥 Pobierz Wykaz", plik, "Wykaz.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                        else: st.error(blad)
                    else: st.warning("Brak uczestników.")

            # Odstęp
            st.write("")
            
            st.subheader("📝 Test Sprawdzający")
            with st.container(border=True):
                if st.button("Generuj Test i Klucz", use_container_width=True):
                    with st.spinner("AI tworzy pytania..."):
                        tresc, klucz = generuj_test_bhp(st.session_state.finalna_tresc)
                        st.session_state.cached_test_content = tresc
                        st.session_state.cached_key_content = klucz
                
                if st.session_state.cached_test_content:
                    col_t1, col_t2 = st.columns(2)
                    with col_t1:
                        ctx_test = {'nazwa_szkolenia': f"Szkolenie: {st.session_state.wybrany_zawod}", 'tresc_testu': st.session_state.cached_test_content}
                        plik_test = generuj_docx_prosty("test_szablon.docx", ctx_test, "Test.docx")
                        if plik_test: st.download_button("📥 Test", plik_test, "Test.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                    with col_t2:
                        if st.session_state.cached_key_content:
                            ctx_klucz = {'klucz_odpowiedzi': st.session_state.cached_key_content}
                            plik_klucz = generuj_docx_prosty("klucz_odpowiedzi_szablon.docx", ctx_klucz, "Klucz.docx")
                            if plik_klucz: st.download_button("📥 Klucz", plik_klucz, "Klucz.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

        # --- PRAWA KOLUMNA: Protokół ---
        with col_right:
            st.subheader("📝 Protokół Egzaminu")
            with st.container(border=True):
                
                # --- ZMIANA: MULTISELECT DLA KOMISJI ---
                st.write("Wybierz członków komisji (kolejność: Przewodniczący, Członek, Członek).")
                wybrana_komisja = st.multiselect(
                    "Skład komisji egzaminacyjnej:",
                    options=opcje_komisji,
                    placeholder="Wybierz 1-3 osoby...",
                    key="komisja_multiselect"
                )
                
                # Oceny
                if uczestnicy_dane_lista:
                    with st.expander("Wpisz Oceny", expanded=False):
                        for i, u in enumerate(uczestnicy_dane_lista):
                            u['ocena'] = st.selectbox(f"{u['imie_nazwisko']}", ["5", "4", "3", "2"], key=f"oc_tab_{i}")
                
                if st.button("Generuj Protokół", use_container_width=True):
                    # LOGIKA AUTOMATYCZNEGO WYBORU
                    komisja_do_przetworzenia = wybrana_komisja if wybrana_komisja else opcje_komisji

                    if not komisja_do_przetworzenia:
                        st.error("Brak członków komisji! Uzupełnij Bazę Kadry.")
                    else:
                        if not wybrana_komisja:
                            st.info("ℹ️ Nie wybrano komisji ręcznie – użyto całej listy z bazy.")

                        komisja_nazwiska = []
                        for osoba in komisja_do_przetworzenia:
                            # Bierzemy tylko Imię i Nazwisko (przed pierwszym przecinkiem)
                            imie_nazwisko = osoba.split(',')[0].strip()
                            komisja_nazwiska.append(imie_nazwisko)
                            
                        context = {
                            'rodzaj_szkolenia': f"Szkolenie BHP: {st.session_state.wybrany_zawod}",
                            'data_egzaminu': data_koniec.strftime("%d.%m.%Y"),
                            'nr_kursu': nr_kursu,
                            'komisja_1_nazwisko': komisja_nazwiska[0] if len(komisja_nazwiska)>0 else "",
                            'komisja_2_nazwisko': komisja_nazwiska[1] if len(komisja_nazwiska)>1 else "",
                            'komisja_3_nazwisko': komisja_nazwiska[2] if len(komisja_nazwiska)>2 else "",
                            'miejsce': miejscowosc,
                            'nazwa_organizatora': st.session_state.zapisana_firma,
                            'data_wystawienia': data_wystawienia.strftime("%d.%m.%Y")
                        }
                        plik, blad = generuj_dokument_z_tabela("protokol_egzaminu_szablon_uproszczony.docx", context, uczestnicy_dane_lista, ['imie_nazwisko', 'ocena', 'uwagi'], index_tabeli=2)
                        if plik: st.download_button("📥 Pobierz Protokół", plik, "Protokol.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                        else: st.error(blad)
# =========================================================
    # SEKCJA POBIERANIA ZBIORCZEGO (ZIP)
    # =========================================================
    st.markdown("---")
    st.subheader("📦 Pobierz wszystko")
    st.info("Wygeneruj komplet dokumentacji jednym kliknięciem.")

    if st.button("Generuj paczkę ZIP ze wszystkimi dokumentami", type="primary", use_container_width=True):
        
        # Sprawdzamy czy mamy kluczowe dane
        if not uczestnicy_dane_lista:
            st.error("Brakuje listy uczestników! Nie można wygenerować kompletu.")
        elif not st.session_state.tematyka_z_godzinami:
            st.error("Brakuje tematyki szkolenia! Wróć do Kroku 1 lub odzyskaj tematykę w zakładce 'Przebieg Szkolenia'.")
        else:
            # Inicjalizacja bufora ZIP w pamięci
            zip_buffer = BytesIO()
            
            try:
                with zipfile.ZipFile(zip_buffer, "w") as zf:
                    
                    # 1. ZAŚWIADCZENIA (Dla każdego uczestnika)
                    for u in uczestnicy_dane_lista:
                        context_cert = {
                            'nazwa_organizatora_szkolenia': st.session_state.zapisana_firma,
                            'imie_nazwisko': u['imie_nazwisko'], 
                            'data_urodzenia': u['data_urodzenia'],
                            'nazwa_szkolenia': f"Szkolenie wstępne BHP: {st.session_state.wybrany_zawod}",
                            'forma_szkolenia': "kurs",
                            'nazwa_organizatora': st.session_state.zapisana_firma,
                            'dzien_rozpoczecia': data_start.strftime("%d.%m.%Y"), 
                            'dzien_zakonczenia': data_koniec.strftime("%d.%m.%Y"),
                            'cel_szkolenia': st.session_state.cel_szkolenia_text, 
                            'miejscowosc_szkolenia': miejscowosc,
                            'data_wystawienia_zaswiadczenia': data_wystawienia.strftime("%d.%m.%Y"),
                            'nr_zaswiadczenia_wg_rejestru': f"{nr_kursu}/{u['index']}"
                        }
                        plik = generuj_docx_prosty("certyfikat_szablon.docx", context_cert, "temp.docx")
                        if plik: zf.writestr(f"Zaswiadczenia/Zaswiadczenie_{u['imie_nazwisko']}.docx", plik.getvalue())

                    # 2. REJESTR ZAŚWIADCZEŃ
                    rejestr_dane = [{'numer': f"{nr_kursu}/{i+1}", 'imie_nazwisko': u['imie_nazwisko'], 'uwagi': ''} for i, u in enumerate(uczestnicy_dane_lista)]
                    context_rej = {
                        'rodzaj_szkolenia': "wstępnego", 'nr_kursu': nr_kursu,
                        'kierownik_nazwisko': kierownik_kursu,
                        'data_wystawienia': data_wystawienia.strftime("%d.%m.%Y"),
                        'nazwa_organizatora': st.session_state.zapisana_firma, 'miejsce': miejscowosc
                    }
                    plik, _ = generuj_dokument_z_tabela("rejestr_zaswiadczen_szablon_uproszczony.docx", context_rej, rejestr_dane, ['numer', 'imie_nazwisko', 'podpis_dummy', 'uwagi'], index_tabeli=2)
                    if plik: zf.writestr("Rejestr_Zaswiadczen.docx", plik.getvalue())

                    # 3. TEMATYKA
                    tematyka = st.session_state.tematyka_z_godzinami
                    total_h = sum(t.get('godziny', 0) for t in tematyka if isinstance(t.get('godziny'), int))
                    tematyka_display = [{"nazwa": t.get('nazwa',''), "godziny": t.get('godziny',0), "praktyka": "0"} for t in tematyka]
                    tematyka_display.append({"nazwa": "RAZEM:", "godziny": total_h, "praktyka": "0"})
                    plik, _ = generuj_dokument_z_tabela("tematyka_szablon_uproszczony.docx", {}, tematyka_display, ['nazwa', 'godziny', 'praktyka'])
                    if plik: zf.writestr("Tematyka_Szkolenia.docx", plik.getvalue())

                    # 4. DZIENNIK ZAJĘĆ
                    zajecia, _ = rozplanuj_zajecia(tematyka, data_start)
                    doc_tpl = DocxTemplate("dziennik_zajec_szablon_uproszczony.docx")
                    doc_tpl.render({'nazwa_organizatora': st.session_state.zapisana_firma})
                    bio = BytesIO()
                    doc_tpl.save(bio)
                    bio.seek(0)
                    doc = Document(bio)
                    if doc.tables:
                        table = doc.tables[0]
                        for i, z in enumerate(zajecia):
                            row = table.add_row().cells
                            if len(row) >= 6:
                                row[0].text = str(i+1); row[1].text = z['data']; row[2].text = str(z['godziny']); row[3].text = z['przedmiot']; row[4].text = z['temat']
                    final_bio = BytesIO()
                    doc.save(final_bio)
                    final_bio.seek(0)
                    zf.writestr("Dziennik_Zajec.docx", final_bio.getvalue())

                    # 5. DZIENNIK LEKCYJNY (ZIP)
                    
                    wybrani_wykladowcy_zip = st.session_state.get("wykladowcy_multiselect", [])
                    
                    baza_wyk_raw = st.session_state.get("baza_wykladowcow_key", "")
                    opcje_bazy_wyk = [x.strip() for x in baza_wyk_raw.splitlines() if x.strip()]
                    
                    # Decyzja: Wybrani czy Wszyscy?
                    finalna_lista_zip = wybrani_wykladowcy_zip if wybrani_wykladowcy_zip else opcje_bazy_wyk
                    
                    if finalna_lista_zip:
                        wykladowcy_lista = []
                        for linia in finalna_lista_zip:
                            parts = [p.strip() for p in linia.split(',', 2)]
                            if len(parts) == 3:
                                wykladowcy_lista.append({'imie_nazwisko': parts[0], 'miejsce_pracy': parts[1], 'funkcja': parts[2], 'przedmiot': 'Szkolenie wstępne BHP', 'godziny_plan': 0, 'godziny_wykonanie': 0})
                        
                        if wykladowcy_lista:
                            # ... reszta logiki przydziału godzin bez zmian ...
                            total_plan = 0
                            for i, temat in enumerate(tematyka):
                                idx = i % len(wykladowcy_lista)
                                h = temat.get('godziny', 0)
                                if isinstance(h, int):
                                    wykladowcy_lista[idx]['godziny_plan'] += h
                                    wykladowcy_lista[idx]['godziny_wykonanie'] += h
                                    total_plan += h
                            wykladowcy_lista.append({'imie_nazwisko': '', 'miejsce_pracy': '', 'funkcja': '', 'przedmiot': 'RAZEM:', 'godziny_plan': total_plan, 'godziny_wykonanie': total_plan})
                            
                            context_lek = {
                                'nazwa_organizatora': st.session_state.zapisana_firma,
                                'dla_kogo': f"Szkolenie dla: {st.session_state.wybrany_zawod}",
                                'data_od': data_start.strftime("%d.%m.%Y"), 'data_do': data_koniec.strftime("%d.%m.%Y"),
                                'miejsce': miejscowosc, 'kierownik_nazwisko': kierownik_kursu, 'kierownik_miejsce_pracy_funkcja': "Kierownik Szkolenia"
                            }
                            plik, _ = generuj_dokument_z_tabela("dziennik_lekcyjny_szablon_uproszczony.docx", context_lek, wykladowcy_lista, ['imie_nazwisko', 'miejsce_pracy', 'funkcja', 'przedmiot', 'godziny_plan', 'godziny_wykonanie'], index_tabeli=4)
                            if plik: zf.writestr("Dziennik_Lekcyjny.docx", plik.getvalue())

                    # 6. PROTOKÓŁ EGZAMINU (ZIP)
                    wybrana_komisja_zip = st.session_state.get("komisja_multiselect", [])
                    baza_kom_raw = st.session_state.get("baza_komisji_key", "")
                    opcje_bazy_kom = [x.strip() for x in baza_kom_raw.splitlines() if x.strip()]
                    
                    finalna_komisja_zip = wybrana_komisja_zip if wybrana_komisja_zip else opcje_bazy_kom

                    if finalna_komisja_zip:
                        komisja_nazwiska = [osoba.split(',')[0].strip() for osoba in finalna_komisja_zip]
                       
                        
                        context_prot = {
                            'rodzaj_szkolenia': f"Szkolenie BHP: {st.session_state.wybrany_zawod}",
                            'data_egzaminu': data_koniec.strftime("%d.%m.%Y"),
                            'nr_kursu': nr_kursu,
                            'komisja_1_nazwisko': komisja_nazwiska[0] if len(komisja_nazwiska)>0 else "",
                            'komisja_2_nazwisko': komisja_nazwiska[1] if len(komisja_nazwiska)>1 else "",
                            'komisja_3_nazwisko': komisja_nazwiska[2] if len(komisja_nazwiska)>2 else "",
                            'miejsce': miejscowosc,
                            'nazwa_organizatora': st.session_state.zapisana_firma,
                            'data_wystawienia': data_wystawienia.strftime("%d.%m.%Y")
                        }
                        plik, _ = generuj_dokument_z_tabela("protokol_egzaminu_szablon_uproszczony.docx", context_prot, uczestnicy_dane_lista, ['imie_nazwisko', 'ocena', 'uwagi'], index_tabeli=2)
                        if plik: zf.writestr("Protokol_Egzaminu.docx", plik.getvalue())

                    # 7. WYKAZ UCZESTNIKÓW
                    plik, _ = generuj_dokument_z_tabela("wykaz_uczestnikow_szablon_uproszczony.docx", {}, uczestnicy_dane_lista, ['imie_nazwisko', 'miejsce_pracy', 'funkcja', 'data_urodzenia'])
                    if plik: zf.writestr("Wykaz_Uczestnikow.docx", plik.getvalue())

                    # 8. TEST I KLUCZ (Opcjonalnie, jeśli wygenerowane)
                    if st.session_state.cached_test_content:
                         ctx_test = {'nazwa_szkolenia': f"Szkolenie: {st.session_state.wybrany_zawod}", 'tresc_testu': st.session_state.cached_test_content}
                         plik = generuj_docx_prosty("test_szablon.docx", ctx_test, "temp.docx")
                         if plik: zf.writestr("Test_BHP.docx", plik.getvalue())
                         
                         if st.session_state.cached_key_content:
                             ctx_klucz = {'klucz_odpowiedzi': st.session_state.cached_key_content}
                             plik = generuj_docx_prosty("klucz_odpowiedzi_szablon.docx", ctx_klucz, "temp.docx")
                             if plik: zf.writestr("Klucz_Odpowiedzi.docx", plik.getvalue())

                    # 9. TREŚĆ SZKOLENIA (TXT)
                    zf.writestr(f"Materialy_Szkoleniowe_{st.session_state.wybrany_zawod}.txt", st.session_state.finalna_tresc)

                # ZAKOŃCZENIE I POBIERANIE
                zip_buffer.seek(0)
                st.success("Paczka dokumentów gotowa!")
                st.download_button(
                    label="📦 POBIERZ PLIK ZIP",
                    data=zip_buffer,
                    file_name=f"Komplet_BHP_{st.session_state.wybrany_zawod}.zip",
                    mime="application/zip",
                    use_container_width=True
                )

            except Exception as e:
                st.error(f"Wystąpił błąd podczas tworzenia archiwum ZIP: {e}")

    st.markdown("---")
    if st.button("🔄 Zacznij od nowa (Nowe Szkolenie)", type="secondary"):
        st.session_state.etap = 1
        st.rerun()